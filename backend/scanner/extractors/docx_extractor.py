"""
Word文档提取器 - 支持 .docx 和 .doc 格式
"""
from pathlib import Path
import subprocess
import tempfile
import platform
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseExtractor
from models.schemas import FileInfo, DocumentMetrics


class DocxExtractor(BaseExtractor):
    """Word文档提取器"""
    
    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def extract(self, file_path: Path, file_info: FileInfo) -> DocumentMetrics:
        """提取Word文档指标"""
        metrics = DocumentMetrics()
        
        # 根据格式选择不同的处理方式
        if file_path.suffix.lower() == '.doc':
            return self._extract_doc(file_path, file_info, metrics)
        else:
            return self._extract_docx(file_path, file_info, metrics)
    
    def _extract_docx(self, file_path: Path, file_info: FileInfo, metrics: DocumentMetrics) -> DocumentMetrics:
        """提取 .docx 格式"""
        try:
            doc = Document(str(file_path))
            
            # 段落统计
            paragraphs = doc.paragraphs
            metrics.paragraph_count = len(paragraphs)
            
            # 字符统计
            total_text = []
            heading_count = 0
            
            for para in paragraphs:
                text = para.text.strip()
                if text:
                    total_text.append(text)
                # 检查是否是标题
                if para.style and para.style.name.startswith('Heading'):
                    heading_count += 1
            
            full_text = '\n'.join(total_text)
            metrics.char_count = len(full_text)
            metrics.word_count = len(full_text.split())
            metrics.heading_count = heading_count
            
            # 表格统计
            metrics.table_count = len(doc.tables)
            
            # 合并单元格统计
            merged_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if hasattr(cell, '_tc'):
                            tc = cell._tc
                            if tc.get('gridSpan') or tc.get('vMerge'):
                                merged_count += 1
            metrics.merged_cell_count = merged_count
            
            # 图片统计
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
            metrics.image_count = image_count
            
        except PackageNotFoundError:
            file_info.is_corrupted = True
            file_info.parse_success = False
            file_info.parse_error = "文件损坏或格式不正确"
        except Exception as e:
            file_info.parse_success = False
            file_info.parse_error = str(e)
        
        return metrics
    
    def _extract_doc(self, file_path: Path, file_info: FileInfo, metrics: DocumentMetrics) -> DocumentMetrics:
        """提取老版 .doc 格式 (使用系统工具)"""
        try:
            text = self._extract_doc_text(file_path)
            
            if text:
                metrics.char_count = len(text)
                metrics.word_count = len(text.split())
                
                # 按换行分段落
                lines = [l for l in text.split('\n') if l.strip()]
                metrics.paragraph_count = len(lines)
                
                # 解析成功
                file_info.parse_success = True
            else:
                # 无法提取文本，但不算失败
                file_info.parse_success = True
                metrics.char_count = 0
                
        except Exception as e:
            # .doc解析失败，标记但不严重
            file_info.parse_success = False
            file_info.parse_error = f"老版.doc格式: {str(e)[:30]}"
        
        return metrics
    
    def _extract_doc_text(self, file_path: Path) -> str:
        """使用系统工具提取.doc文本"""
        system = platform.system()
        
        if system == "Darwin":  # macOS
            # 使用 textutil 转换为纯文本
            try:
                result = subprocess.run(
                    ["textutil", "-convert", "txt", "-stdout", str(file_path)],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
        
        # 其他平台或textutil失败，尝试 antiword
        try:
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        # 都失败了，返回空
        return ""
    
    def extract_text(self, file_path: Path) -> str:
        """提取文本内容"""
        if file_path.suffix.lower() == '.doc':
            return self._extract_doc_text(file_path)
        
        try:
            doc = Document(str(file_path))
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            return '\n'.join(texts)
        except Exception:
            return ""

